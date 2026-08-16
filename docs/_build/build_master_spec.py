from __future__ import annotations

import re
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
SOURCE = ROOT / "SOLUTION_SPEC.md"
OUTPUT = ROOT / "Adaptive_Payment_Assurance_Range_Spec.docx"
ASSET_DIR = ROOT / "docs" / "_build" / "assets"

NAVY = "0B2545"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MID_BLUE = "2A6F97"
PALE_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"
MID_GRAY = "667085"
INK = "17202A"
ORANGE = "C26A1B"
WHITE = "FFFFFF"


def rgb(value: str) -> RGBColor:
    return RGBColor.from_string(value)


def set_run_font(run, *, name="Calibri", size=None, color=INK, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_fixed_layout(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths[index] / 1440)
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_row_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def create_numbering(doc: Document, kind: str, start_value: int = 1) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [
        int(node.get(qn("w:numId")))
        for node in numbering.findall(qn("w:num"))
    ]
    abstract_id = max(abstract_ids, default=-1) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), str(start_value))
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if kind == "bullet" else "decimal")
    level.append(num_fmt)
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "•" if kind == "bullet" else "%1.")
    level.append(level_text)
    suffix = OxmlElement("w:suff")
    suffix.set(qn("w:val"), "space")
    level.append(suffix)
    p_pr = OxmlElement("w:pPr")
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "540")
    indent.set(qn("w:hanging"), "270")
    p_pr.append(indent)
    level.append(p_pr)
    if kind == "bullet":
        r_pr = OxmlElement("w:rPr")
        fonts = OxmlElement("w:rFonts")
        fonts.set(qn("w:ascii"), "Arial")
        fonts.set(qn("w:hAnsi"), "Arial")
        r_pr.append(fonts)
        level.append(r_pr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id: int):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num])
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run_font(run, size=9, color=MID_GRAY)
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr, separate, text, end])


def configure_document(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    code_style = styles.add_style("Code Block", 1)
    code_style.font.name = "Menlo"
    code_style._element.rPr.rFonts.set(qn("w:ascii"), "Menlo")
    code_style._element.rPr.rFonts.set(qn("w:hAnsi"), "Menlo")
    code_style.font.size = Pt(8.5)
    code_style.font.color.rgb = rgb(NAVY)
    code_style.paragraph_format.left_indent = Inches(0.18)
    code_style.paragraph_format.right_indent = Inches(0.18)
    code_style.paragraph_format.space_before = Pt(4)
    code_style.paragraph_format.space_after = Pt(7)
    code_style.paragraph_format.line_spacing = 1.05

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_run = header.add_run("Adaptive Payment Assurance Range | Solution specification")
    set_run_font(header_run, size=8.5, color=MID_GRAY, bold=True)
    add_page_field(section.footer.paragraphs[0])


def add_cover(doc: Document):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(76)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("ADAPTIVE PAYMENT\nASSURANCE RANGE")
    set_run_font(run, size=28, color=NAVY, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(28)
    run = p.add_run("Solution specification for the Mastercard Innovation Challenge 2026")
    set_run_font(run, size=14, color=DARK_BLUE)

    metadata = [
        ("Purpose", "Evidence-backed adversarial assurance for card, A2A, and agentic-commerce controls"),
        ("Status", "Approved target architecture"),
        ("Version", "1.1"),
        ("Date", "16 August 2026"),
    ]
    for label, value in metadata:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        label_run = p.add_run(f"{label}: ")
        set_run_font(label_run, size=10.5, color=MID_GRAY, bold=True)
        value_run = p.add_run(value)
        set_run_font(value_run, size=10.5, color=INK)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(32)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("DECISION")
    set_run_font(run, size=10, color=ORANGE, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(
        "Build a pre-production assurance range that converts sourced GenAI payment threats into "
        "constrained synthetic campaigns, challenges layered defenses under hidden shifts, and "
        "produces an auditable human promotion decision."
    )
    set_run_font(run, size=12, color=NAVY, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(28)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("Editable diagrams and subsystem specifications are included in the repository under docs/.")
    set_run_font(run, size=9.5, color=MID_GRAY, italic=True)

    doc.add_page_break()


def clean_inline(text: str) -> str:
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", text)
    text = text.replace("**", "").replace("`", "")
    text = text.replace("  ", " ")
    return text.strip()


def add_paragraph_with_inline(doc, text: str, style=None, num_id=None):
    p = doc.add_paragraph(style=style)
    parts = re.split(r"(\*\*.*?\*\*|`.*?`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            set_run_font(run, bold=True)
        elif part.startswith("`") and part.endswith("`"):
            run = p.add_run(part[1:-1])
            set_run_font(run, name="Menlo", size=9, color=NAVY)
        else:
            linked = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", part)
            run = p.add_run(linked)
            set_run_font(run)
    if num_id is not None:
        apply_numbering(p, num_id)
    return p


def add_table(doc, rows):
    if not rows:
        return
    col_count = max(len(row) for row in rows)
    if col_count == 2:
        widths = [2700, 6660]
    elif col_count == 3:
        widths = [2100, 3360, 3900]
    elif col_count == 4:
        widths = [1750, 2350, 2470, 2790]
    else:
        base = 9360 // col_count
        widths = [base] * col_count
        widths[-1] += 9360 - sum(widths)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"
    set_table_fixed_layout(table, widths)
    for row_index, row in enumerate(rows):
        for col_index in range(col_count):
            cell = table.cell(row_index, col_index)
            text = clean_inline(row[col_index] if col_index < len(row) else "")
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.08
            run = p.add_run(text)
            set_run_font(run, size=9.2, color=INK, bold=row_index == 0)
            if row_index == 0:
                set_cell_shading(cell, PALE_BLUE)
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    for row in table.rows:
        set_row_cant_split(row)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def add_module_table(doc, code_lines):
    rows = [["Module", "Responsibility"]]
    for raw in code_lines:
        line = raw.strip()
        if not line or line == "src/":
            continue
        match = re.match(r"([^\s]+)\s{2,}(.+)", line)
        if match:
            rows.append([f"src/{match.group(1)}", match.group(2)])
        else:
            rows.append([f"src/{line}", ""])
    add_table(doc, rows)


def add_contents(doc: Document):
    p = doc.add_paragraph("Contents", style="Heading 1")
    p.paragraph_format.space_before = Pt(0)
    sections = [
        "Executive decision and product thesis",
        "Goals, users, and architecture",
        "Threat portfolio and scenario requirements",
        "Adaptive red-team and layered defense",
        "Agentic trust and data architecture",
        "Evaluation, safety, prototype, and operations",
        "Acceptance criteria, risks, documentation, and approval",
    ]
    for section in sections:
        if "contents_num_id" not in locals():
            contents_num_id = create_numbering(doc, "bullet")
        add_paragraph_with_inline(doc, section, num_id=contents_num_id)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run("The detailed repository documentation and editable Mermaid figures remain the source of truth for implementation-level contracts.")
    set_run_font(run, size=10, color=MID_GRAY, italic=True)
    doc.add_page_break()


def add_diagram_reference(doc, number: int):
    mapping = {
        1: ("System context", "docs/diagrams/01-system-context.mmd"),
        2: ("Identify, Generate, Defend, and Assure architecture", "docs/diagrams/02-logical-architecture.mmd"),
        3: ("Payment lifecycle", "docs/diagrams/05-payment-lifecycle.mmd"),
    }
    title, path = mapping.get(number, ("Architecture diagram", "docs/diagrams/README.md"))
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_fixed_layout(table, [9360])
    cell = table.cell(0, 0)
    set_row_cant_split(table.rows[0])
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    set_cell_shading(cell, LIGHT_GRAY)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(f"Figure {number}. {title}\nEditable source: {path}")
    set_run_font(run, size=10, color=NAVY, bold=True)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def parse_markdown(doc: Document, source: str):
    lines = source.splitlines()
    index = 0
    table_rows = []
    in_code = False
    code_lang = ""
    code_lines = []
    mermaid_count = 0
    skipped_opening = 0
    current_list_kind = None
    current_list_num_id = None

    def flush_table():
        nonlocal table_rows
        if table_rows:
            filtered = [row for row in table_rows if not all(re.fullmatch(r":?-{3,}:?", c.strip()) for c in row)]
            add_table(doc, filtered)
            table_rows = []

    while index < len(lines):
        raw = lines[index]
        line = raw.rstrip()

        if line.startswith("```"):
            flush_table()
            if not in_code:
                in_code = True
                code_lang = line[3:].strip()
                code_lines = []
            else:
                if code_lang == "mermaid":
                    mermaid_count += 1
                    add_diagram_reference(doc, mermaid_count)
                elif code_lang == "text" and code_lines and code_lines[0].strip() == "src/":
                    add_module_table(doc, code_lines)
                else:
                    p = doc.add_paragraph(style="Code Block")
                    run = p.add_run("\n".join(code_lines))
                    set_run_font(run, name="Menlo", size=8.5, color=NAVY)
                in_code = False
                code_lang = ""
                code_lines = []
            index += 1
            continue

        if in_code:
            code_lines.append(raw)
            index += 1
            continue

        if line.startswith("|") and line.endswith("|"):
            cells = [cell.strip() for cell in line.strip("|").split("|")]
            table_rows.append(cells)
            index += 1
            continue
        flush_table()

        if not line.strip():
            current_list_kind = None
            current_list_num_id = None
            index += 1
            continue

        if line.startswith("# ") or line.startswith("## Solution specification"):
            current_list_kind = None
            current_list_num_id = None
            skipped_opening += 1
            index += 1
            continue
        if line.startswith("## "):
            current_list_kind = None
            current_list_num_id = None
            doc.add_paragraph(clean_inline(line[3:]), style="Heading 1")
        elif line.startswith("### "):
            current_list_kind = None
            current_list_num_id = None
            doc.add_paragraph(clean_inline(line[4:]), style="Heading 2")
        elif line.startswith("#### "):
            current_list_kind = None
            current_list_num_id = None
            doc.add_paragraph(clean_inline(line[5:]), style="Heading 3")
        elif re.match(r"^\d+\.\s+", line):
            marker = int(re.match(r"^(\d+)\.\s+", line).group(1))
            text = re.sub(r"^\d+\.\s+", "", line)
            if current_list_kind != "decimal":
                current_list_num_id = create_numbering(doc, "decimal", marker)
                current_list_kind = "decimal"
            add_paragraph_with_inline(doc, text, num_id=current_list_num_id)
        elif line.startswith("- "):
            if current_list_kind != "bullet":
                current_list_num_id = create_numbering(doc, "bullet")
                current_list_kind = "bullet"
            add_paragraph_with_inline(doc, line[2:], num_id=current_list_num_id)
        elif line.startswith("> "):
            current_list_kind = None
            current_list_num_id = None
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.22)
            p.paragraph_format.right_indent = Inches(0.22)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(10)
            run = p.add_run(clean_inline(line[2:]))
            set_run_font(run, size=11.5, color=NAVY, bold=True)
        elif re.match(r"^\*\*[^*]+:\*\*", line):
            current_list_kind = None
            current_list_num_id = None
            add_paragraph_with_inline(doc, line)
        else:
            current_list_kind = None
            current_list_num_id = None
            add_paragraph_with_inline(doc, line)
        index += 1
    flush_table()


def build():
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_document(doc)
    add_cover(doc)
    add_contents(doc)
    parse_markdown(doc, SOURCE.read_text(encoding="utf-8"))
    props = doc.core_properties
    props.title = "Adaptive Payment Assurance Range - Solution Specification"
    props.subject = "Mastercard Innovation Challenge 2026 proposed solution"
    props.author = "Mastercard Innovation Challenge Team"
    props.keywords = "payment fraud, GenAI, red teaming, assurance, agentic commerce"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
