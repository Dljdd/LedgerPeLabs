#!/usr/bin/env python3
"""Build the final Kaggle solution walkthrough for team LedgerPeLabs."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips


ROOT = Path(__file__).resolve().parents[2]
OUTPUT = ROOT / "docs" / "submission" / "LedgerPeLabs.docx"


INK = "19150F"
INK_SOFT = "3F382E"
MUTED = "6B5F50"
CREAM = "F1E9DC"
CREAM_LIGHT = "F8F4ED"
ORANGE = "E8793D"
ORANGE_LIGHT = "F6E5D8"
GREEN = "4D8A66"
LINE = "D8CFC2"
WHITE = "FFFFFF"
DEFAULT_CELL_MARGINS_DXA = {"top": 80, "bottom": 80, "start": 120, "end": 120}


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def ensure_child(parent, tag: str):
    child = parent.find(qn(tag))
    if child is None:
        child = OxmlElement(tag)
        parent.append(child)
    return child


def set_width(parent, tag: str, width_dxa: int) -> None:
    width = ensure_child(parent, tag)
    width.set(qn("w:type"), "dxa")
    width.set(qn("w:w"), str(width_dxa))


def apply_table_geometry(
    table,
    column_widths_dxa: list[int],
    *,
    indent_dxa: int | None = None,
    cell_margins_dxa: dict[str, int] | None = None,
) -> None:
    """Synchronize Word table, grid, and cell widths without local helpers."""

    widths = [int(width) for width in column_widths_dxa]
    if not widths or any(width <= 0 for width in widths):
        raise ValueError("column widths must be positive")

    margins = dict(DEFAULT_CELL_MARGINS_DXA)
    if cell_margins_dxa:
        margins.update({key: int(value) for key, value in cell_margins_dxa.items()})

    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table_properties = table._tbl.tblPr
    set_width(table_properties, "w:tblW", sum(widths))

    table_indent = ensure_child(table_properties, "w:tblInd")
    table_indent.set(qn("w:type"), "dxa")
    table_indent.set(qn("w:w"), str(margins["start"] if indent_dxa is None else indent_dxa))
    ensure_child(table_properties, "w:tblLayout").set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_column = OxmlElement("w:gridCol")
        grid_column.set(qn("w:w"), str(width))
        grid.append(grid_column)

    for column_index, width in enumerate(widths):
        table.columns[column_index].width = Twips(width)
    for row in table.rows:
        if len(row.cells) != len(widths):
            raise ValueError("table rows must be unmerged before geometry is applied")
        for column_index, cell in enumerate(row.cells):
            width = widths[column_index]
            cell.width = Twips(width)
            cell_properties = cell._tc.get_or_add_tcPr()
            set_width(cell_properties, "w:tcW", width)
            cell_margins = ensure_child(cell_properties, "w:tcMar")
            for side in ("top", "bottom", "start", "end"):
                margin = ensure_child(cell_margins, f"w:{side}")
                margin.set(qn("w:w"), str(margins[side]))
                margin.set(qn("w:type"), "dxa")


def set_run_font(
    run,
    *,
    name: str = "Calibri",
    size: float | None = None,
    color: str = INK,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_picture_alt(inline_shape, alt_text: str) -> None:
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("descr", alt_text)
    doc_pr.set("title", alt_text)


def set_cell_fill(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def set_cell_borders(cell, color: str = LINE, size: str = "6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:color"), color)


def set_keep_with_next(paragraph) -> None:
    paragraph.paragraph_format.keep_with_next = True


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run_font(run, size=8.5, color=MUTED)
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_1)
    run._r.append(instr)
    run._r.append(fld_char_2)


def add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    relation_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relation_id)
    new_run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), ORANGE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_properties.append(color)
    run_properties.append(underline)
    new_run.append(run_properties)
    text_node = OxmlElement("w:t")
    text_node.text = text
    new_run.append(text_node)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    heading_tokens = {
        "Heading 1": (16, INK, 16, 8),
        "Heading 2": (13, ORANGE, 12, 6),
        "Heading 3": (12, INK_SOFT, 8, 4),
    }
    for style_name, (size, color, before, after) in heading_tokens.items():
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("APAR")
    set_run_font(run, name="Georgia", size=9, color=INK, bold=True)
    run = p.add_run("   |   MASTERCARD INNOVATION CHALLENGE 2026")
    set_run_font(run, size=8.5, color=MUTED, bold=True)

    footer = section.footer
    table = footer.add_table(rows=1, cols=2, width=Inches(6.5))
    table.autofit = False
    apply_table_geometry(table, [7200, 2160], indent_dxa=0)
    left = table.cell(0, 0).paragraphs[0]
    left.paragraph_format.space_after = Pt(0)
    run = left.add_run("LedgerPeLabs · Solution walkthrough")
    set_run_font(run, size=8.5, color=MUTED)
    right = table.cell(0, 1).paragraphs[0]
    right.paragraph_format.space_after = Pt(0)
    add_page_number(right)
    for cell in table.rows[0].cells:
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_borders(cell, WHITE, "0")

    props = doc.core_properties
    props.title = "APAR - Adaptive Payment Assurance Range"
    props.subject = "Mastercard Innovation Challenge 2026 solution walkthrough"
    props.author = "Dylan Moraes"
    props.keywords = "payment fraud, red teaming, simulation, graph features, agentic payments"


def add_title_block(doc: Document) -> None:
    cover = ROOT / "docs" / "demo" / "media" / "apar-project-cover-master.png"
    picture = doc.add_picture(str(cover), width=Inches(6.5))
    set_picture_alt(
        picture,
        "APAR project cover showing the Identify, Generate, Defend, and Assure loop.",
    )
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.paragraphs[-1].paragraph_format.space_after = Pt(24)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("SOLUTION WALKTHROUGH")
    set_run_font(run, size=10, color=ORANGE, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("APAR")
    set_run_font(run, name="Georgia", size=30, color=INK, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(20)
    run = p.add_run("Adaptive Payment Assurance Range")
    set_run_font(run, size=16, color=INK_SOFT, bold=True)

    add_callout(
        doc,
        "Assure the campaign. Verify the intent. Promote only the evidence.",
        "APAR is a pre-production assurance range that turns emerging GenAI-enabled "
        "payment threats into bounded synthetic campaigns, executes them through "
        "rail-correct lifecycles, compares defenses, verifies delegated intent, and "
        "preserves a human-controlled promotion boundary.",
    )

    metadata = doc.add_table(rows=4, cols=2)
    rows = [
        ("Team", "LedgerPeLabs"),
        ("Track", "AI Defense Lab for Payment Security"),
        ("Competition", "Mastercard Innovation Challenge 2026"),
        ("Submission date", "31 August 2026"),
    ]
    for index, (label, value) in enumerate(rows):
        metadata.cell(index, 0).text = label
        metadata.cell(index, 1).text = value
    apply_table_geometry(metadata, [2700, 6660])
    style_table(metadata, header=False, label_column=True)
    doc.add_page_break()


def add_callout(doc: Document, label: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_fill(cell, ORANGE_LIGHT)
    set_cell_borders(cell, ORANGE, "10")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(label)
    set_run_font(run, size=11, color=INK, bold=True)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(body)
    set_run_font(run, size=10.5, color=INK_SOFT)
    apply_table_geometry(
        table,
        [9360],
        cell_margins_dxa={"top": 180, "bottom": 180, "start": 220, "end": 220},
    )
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def style_table(table, *, header: bool = True, label_column: bool = False) -> None:
    if header:
        tr_pr = table.rows[0]._tr.get_or_add_trPr()
        tbl_header = OxmlElement("w:tblHeader")
        tbl_header.set(qn("w:val"), "true")
        tr_pr.append(tbl_header)
    for row_index, row in enumerate(table.rows):
        for cell_index, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_borders(cell)
            if header and row_index == 0:
                set_cell_fill(cell, INK)
            elif row_index % 2 == 0:
                set_cell_fill(cell, CREAM_LIGHT)
            else:
                set_cell_fill(cell, WHITE)
            if label_column and cell_index == 0:
                set_cell_fill(cell, CREAM_LIGHT)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.08
                for run in paragraph.runs:
                    if header and row_index == 0:
                        set_run_font(run, size=9.2, color=CREAM, bold=True)
                    elif label_column and cell_index == 0:
                        set_run_font(run, size=9.5, color=INK, bold=True)
                    else:
                        set_run_font(run, size=9.4, color=INK)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]):
    table = doc.add_table(rows=1, cols=len(headers))
    for index, header in enumerate(headers):
        table.cell(0, index).text = header
    for row_data in rows:
        row = table.add_row()
        for index, value in enumerate(row_data):
            row.cells[index].text = value
    apply_table_geometry(table, widths)
    style_table(table)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)


def add_number(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)


def add_image(
    doc: Document,
    relative_path: str,
    caption: str,
    *,
    width_inches: float = 5.2,
    dedicated_page: bool = True,
) -> None:
    path = ROOT / relative_path
    if dedicated_page:
        doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    picture = run.add_picture(str(path), width=Inches(width_inches))
    set_picture_alt(picture, caption)
    caption_p = doc.add_paragraph()
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_p.paragraph_format.space_before = Pt(2)
    caption_p.paragraph_format.space_after = Pt(10)
    run = caption_p.add_run(caption)
    set_run_font(run, size=8.5, color=MUTED, italic=True)


def add_link_line(doc: Document, label: str, url: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f"{label}: ")
    set_run_font(run, size=10.5, color=INK, bold=True)
    add_hyperlink(p, url, url)


def build() -> None:
    doc = Document()
    configure_document(doc)
    add_title_block(doc)

    doc.add_heading("1. Executive summary", level=1)
    doc.add_paragraph(
        "Generative AI does not create a new payment rail; it changes attacker speed, "
        "personalization, coordination, and autonomy. APAR therefore evaluates fraud as "
        "a campaign and a control-system problem rather than as isolated transaction "
        "classification. The system closes the loop from threat evidence to simulation, "
        "defense comparison, failure discovery, and governed promotion evidence."
    )
    add_callout(
        doc,
        "Competition fit",
        "APAR directly implements Identify, Generate, and Defend, then adds an assurance "
        "layer so the evidence is causal, reproducible, tamper-evident, and suitable for "
        "human review.",
    )
    add_table(
        doc,
        ["Evaluation area", "APAR response", "Demonstrable proof"],
        [
            ["Attack diversity", "Evidence-backed threat registry", "20 typed threat cards; four executable families"],
            ["Simulation fidelity", "Stateful rail and ledger execution", "Card, A2A, and agentic adapters; lifecycle and conservation tests"],
            ["Detection efficacy", "Calibrated graph-aware ensemble", "Four-arm comparison and exact portable-model replay"],
            ["Novelty", "Campaign assurance plus intent integrity", "Adaptive search, TrustVerifier, and content-addressed evidence"],
            ["Feasibility", "Shadow-mode adoption path", "CPU scorer, selective actions, capacity gates, and human promotion"],
        ],
        [1900, 3400, 4060],
    )

    doc.add_heading("Team LedgerPeLabs", level=2)
    doc.add_paragraph(
        "The competition team brings together product, engineering, fraud-defense, and "
        "research perspectives."
    )
    for team_member in (
        "Dylan Moraes",
        "Anuj Sharma",
        "Dhananjay Joshi",
        "Rahul Biradar",
    ):
        add_bullet(doc, team_member)

    doc.add_page_break()
    doc.add_heading("2. Identify: emerging GenAI-enabled payment threats", level=1)
    doc.add_paragraph(
        "APAR separates observed threat evidence from modeled GenAI capability changes. "
        "Each threat card records the affected rail, attack stages, observables, "
        "confidence, provenance, simulation support, and safety classification. This "
        "prevents an LLM-generated idea from silently becoming a factual claim."
    )
    doc.add_paragraph("The implemented deep-dive portfolio contains four campaign families:")
    add_bullet(doc, "AI-personalized APP scams with mule-network convergence and staged cash-out.")
    add_bullet(doc, "Adaptive card-testing and card-not-present bursts that vary pacing, amounts, and reuse.")
    add_bullet(doc, "Synthetic merchant, identity, refund, dispute, and payout abuse across lifecycle stages.")
    add_bullet(doc, "Agentic-payment intent abuse: mandate violation, cart mutation, merchant substitution, replay, and token-scope abuse.")
    doc.add_paragraph(
        "The broader registry also covers credential compromise, deepfake-assisted social "
        "engineering, multilingual personalization, onboarding abuse, cross-channel "
        "coordination, and autonomous reconnaissance. Executable depth is deliberately "
        "bounded so every scenario remains reviewable and safe."
    )
    add_table(
        doc,
        ["Executable family", "Payment context", "Campaign structure"],
        [
            ["APP scam and mule", "A2A", "Victim persuasion -> authorized transfer -> convergence -> cash-out"],
            ["Card testing CNP", "Card", "Low-value probes -> retry adaptation -> merchant/device burst"],
            ["Synthetic merchant/refund", "Card", "Onboarding -> authorization -> refund/dispute -> payout"],
            ["Agentic intent abuse", "Agentic", "Mandate or cart mutation -> authority verification -> payment attempt"],
        ],
        [2300, 1800, 5260],
    )
    add_image(
        doc,
        "docs/demo/screenshots/apar-console-overview-desktop.png",
        "APAR Overview: sourced threat framing and the closed Identify-Generate-Defend-Assure loop.",
    )

    heading = doc.add_heading("3. Generate: high-fidelity synthetic campaigns", level=1)
    heading.paragraph_format.page_break_before = True
    doc.add_paragraph(
        "Fraud cases are not hand-authored rows. The production corpus path runs "
        "PopulationGenerator -> CampaignGenerator -> SimulationEngine -> rail adapter -> "
        "events -> ledger -> validated evidence -> decision rows. The event-to-row "
        "projection preserves concrete event IDs, amounts, rails, lifecycle states, and "
        "decision-time availability."
    )
    doc.add_heading("Fidelity controls", level=2)
    add_bullet(doc, "Explicit authorization, transfer, settlement, refund, reversal, and recovery states.")
    add_bullet(doc, "Separate event, ingestion, feature, and decision timestamps; future data is excluded.")
    add_bullet(doc, "Ledger conservation and lifecycle dependencies; failed or reversed value cannot fund cash-out.")
    add_bullet(doc, "Benign novelty including retries, shared devices, travel, missing fields, and changing channel use.")
    add_bullet(doc, "Bounded attacker budgets, declared feedback, hidden validity constraints, and deterministic replay manifests.")
    doc.add_paragraph(
        "The generator is useful because it produces campaign structure and operational "
        "edge cases, not just statistically similar independent rows. The resulting data "
        "supports both model training and stress testing."
    )
    add_table(
        doc,
        ["Simulator contract", "Preserved evidence", "Why it matters"],
        [
            ["Causal time", "Event, ingestion, feature, and decision clocks", "Prevents future information from reaching a decision"],
            ["Rail lifecycle", "Authorization, transfer, settlement, refund, reversal", "Matches payment-state ownership and recovery semantics"],
            ["Economic validity", "Opening balances, postings, dependencies", "Prevents impossible cash-out and value creation"],
            ["Campaign identity", "Actors, counterparties, edges, source event IDs", "Supports graph features and linked investigation cases"],
            ["Replay", "Seed, manifest, config, hashes", "Makes scenarios deterministic and tamper-evident"],
        ],
        [2200, 3400, 3760],
    )
    add_image(
        doc,
        "docs/demo/screenshots/apar-console-replay-desktop.png",
        "Replay route: synchronized payment lifecycle, campaign evidence, and decision trace.",
        width_inches=4.7,
    )

    heading = doc.add_heading("4. Defend: campaign-aware detection and mitigation", level=1)
    heading.paragraph_format.page_break_before = True
    doc.add_paragraph(
        "The selected competition model is a portable three-member calibrated CatBoost "
        "ensemble over 46 frozen, past-only features. It combines transaction context "
        "with actor and counterparty velocity, amount deviation, prior pair behavior, "
        "fan-in/fan-out, shared neighbors, two-hop reach, burst motifs, component and "
        "density summaries, lifecycle signals, and data-quality flags."
    )
    doc.add_paragraph(
        "The action policy is selective rather than binary: approve, challenge, "
        "review-hold, or decline-hold. Selection therefore considers customer friction, "
        "review capacity, latency, calibration, and captured value alongside recall and "
        "precision."
    )
    add_table(
        doc,
        ["Arm", "Recall", "Precision", "F1", "False decline", "Challenge", "p95 ms"],
        [
            ["Rules only", "85.949%", "14.878%", "25.366%", "87.436%", "2.639%", "4.062"],
            ["Ensemble, no graph", "99.745%", "94.756%", "97.186%", "0.000%", "0.802%", "3.382"],
            ["Ensemble with graph", "99.867%", "95.876%", "97.831%", "0.0037%", "0.572%", "3.544"],
            ["Full Sentinel hybrid", "99.929%", "16.846%", "28.831%", "87.437%", "2.802%", "19.742"],
        ],
        [1900, 1120, 1180, 1020, 1450, 1370, 1320],
    )
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("Evidence boundary. ")
    set_run_font(run, size=9.5, color=INK, bold=True)
    run = p.add_run(
        "These are verified recovered diagnostics on a synthetic development corpus. "
        "They support architecture comparison; they are not official Stage 70, "
        "production, real-cardholder, Mastercard-data, or external-validation results."
    )
    set_run_font(run, size=9.5, color=MUTED, italic=True)

    doc.add_heading("5. What the experiments changed", level=1)
    add_number(doc, "Rules were too blunt: fraud coverage came with unacceptable benign friction.")
    add_number(doc, "The calibrated non-graph ensemble established a strong real-time baseline.")
    add_number(doc, "An early perfect result exposed future graph leakage; the path was rejected and causal append/equal-time isolation tests were added.")
    add_number(doc, "Constructed fraud rows were replaced by real campaign, simulator, rail, event, and ledger execution.")
    add_number(doc, "Causal graph summaries produced the best precision/recall/friction balance without placing a heavyweight GNN in the synchronous path.")
    doc.add_paragraph(
        "The graph arm improves the already strong non-graph ensemble by 0.122 percentage "
        "points of recall, 1.120 points of precision, 0.645 points of F1, and 0.230 points "
        "lower challenge rate for roughly 0.162 ms of additional p95 latency. The full "
        "hybrid has the highest recall but fails the friction gates, which is why it is not "
        "called the champion."
    )
    add_image(
        doc,
        "docs/demo/screenshots/apar-console-investigation-desktop.png",
        "Investigation route: linked campaign context replaces disconnected transaction alerts.",
    )

    heading = doc.add_heading("6. Verifiable intent for agentic payments", level=1)
    heading.paragraph_format.page_break_before = True
    doc.add_paragraph(
        "Agentic commerce introduces an authorization-integrity problem that a fraud "
        "score alone cannot solve. Before statistical risk is considered, APAR's "
        "deterministic TrustVerifier checks the authority chain for the exact purchase."
    )
    add_bullet(doc, "Registered agent identity and signing key.")
    add_bullet(doc, "User mandate, permitted scope, amount, currency, merchant, and category constraints.")
    add_bullet(doc, "Merchant and cart binding plus payment-intent hash.")
    add_bullet(doc, "Token scope, issued-at time, expiry, nonce, and replay state.")
    add_bullet(doc, "Execution receipt and outcome binding.")
    add_callout(
        doc,
        "Deterministic checks precede ML risk",
        "Invalid signatures, expired mandates, scope mismatches, cart mutations, merchant "
        "substitution, and replay fail closed before the model can approve a payment.",
    )

    doc.add_heading("7. Evidence governance and assurance", level=1)
    doc.add_paragraph(
        "Models, scenarios, controls, results, and replay traces are content-addressed. "
        "Independent verifiers recompute hashes and evidence semantics; tampering, label "
        "leakage, future information, and broken rail or ledger provenance fail closed. "
        "Rejected experiments remain visible, and no model can promote itself."
    )
    add_bullet(doc, "Label-shuffle and single-class controls.")
    add_bullet(doc, "Identity-renaming invariance and feature/label leakage tests.")
    add_bullet(doc, "Future-append and equal-time isolation tests.")
    add_bullet(doc, "Rail, lifecycle, ledger, trust, and evidence-tamper tests.")
    add_bullet(doc, "Hash-bound portable bundle and exact probability/action replay.")
    add_bullet(doc, "Explicit human promotion decision with rollback and evidence boundaries.")
    add_image(
        doc,
        "docs/demo/screenshots/apar-console-assurance-focus.png",
        "Assurance route: compact evidence inspection, trust checks, and human promotion boundary.",
        width_inches=2.95,
    )

    heading = doc.add_heading("8. Working web prototype", level=1)
    heading.paragraph_format.page_break_before = True
    doc.add_paragraph(
        "The public six-route console demonstrates the complete loop: threat framing, "
        "scenario construction, synchronized replay, connected investigation, defense "
        "comparison, and assurance evidence. The hosted build uses a committed, "
        "hash-bound trace so the walkthrough remains accessible without a Python service."
    )
    doc.add_paragraph(
        "The repository also contains the real portable scorer. The local console loads "
        "the packaged graph ensemble, scores 12 curated cases, verifies probabilities and "
        "actions exactly, and exposes the results through the same interface."
    )
    add_link_line(doc, "Prototype", "https://web-six-tau-bxhm7rwrzu.vercel.app/overview")
    add_link_line(doc, "Repository", "https://github.com/Dljdd/LedgerPeLabs")
    add_link_line(doc, "Demo video", "https://www.youtube.com/watch?v=A_4Pe_A7iMg")
    doc.add_heading("Local run", level=2)
    for command in (
        "python3.12 -m venv .venv",
        ".venv/bin/python -m pip install -e '.[dev]'",
        "npm ci --prefix web",
        ".venv/bin/python scripts/run_apar_console.py start",
    ):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(command)
        set_run_font(run, name="Courier New", size=9.5, color=INK)

    doc.add_heading("9. Real-world feasibility", level=1)
    doc.add_paragraph(
        "APAR is designed as an assurance layer around existing payment-risk and "
        "fraud-operations systems, not as a replacement authorization switch. A practical "
        "adoption path keeps the model advisory until its evidence and operating impact "
        "are independently validated."
    )
    add_table(
        doc,
        ["Stage", "Use", "Control boundary"],
        [
            ["Offline replay", "Historical and synthetic evaluation", "No customer impact"],
            ["Shadow mode", "Live feature generation and drift measurement", "Scores are advisory only"],
            ["Assisted operations", "Human-reviewed cases and selective challenges", "Capacity and rollback limits"],
            ["Limited challenger", "Small governed traffic slice", "Explicit owner and stop criteria"],
            ["Governed scale", "Continuous adversarial regression", "Independent validation and human promotion"],
        ],
        [1800, 3500, 4060],
    )
    doc.add_paragraph(
        "The synchronous path contains schema checks, agentic integrity, feature lookup, "
        "rules, CPU scoring, and reason codes. Graph expansion, campaign discovery, case "
        "formation, drift analysis, and retraining stay asynchronous. This preserves a "
        "credible path to low-latency live payment integration."
    )

    doc.add_heading("10. Reproducibility, safety, and limitations", level=1)
    add_bullet(doc, "Synthetic data only; no real cardholder data is included.")
    add_bullet(doc, "The 12 visible cases are curated replay demonstrations, not a population estimate.")
    add_bullet(doc, "Recovered four-arm results are verified diagnostics, not official Stage 70 capacity evidence.")
    add_bullet(doc, "The official chain is incomplete at Stage 70; APAR makes no production-readiness claim.")
    add_bullet(doc, "The prototype does not export operational attack instructions or autonomously deploy controls.")
    add_bullet(doc, "Production adoption would require independent data validation, privacy review, monitoring, and model-risk approval.")
    doc.add_paragraph(
        "This boundary is intentional. A payment-assurance system should explain both "
        "what its evidence proves and what it does not. APAR's central contribution is a "
        "closed, testable, and governable process for adapting defenses as attackers "
        "adapt—without granting the model authority over deployment."
    )

    doc.add_heading("Submission links", level=1)
    add_link_line(doc, "Working prototype", "https://web-six-tau-bxhm7rwrzu.vercel.app/overview")
    add_link_line(doc, "Public source repository", "https://github.com/Dljdd/LedgerPeLabs")
    add_link_line(
        doc,
        "Competition deck",
        "https://github.com/Dljdd/LedgerPeLabs/blob/"
        "codex/apar-final-submission/docs/submission/APAR_COMPETITION_DECK.pdf",
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
